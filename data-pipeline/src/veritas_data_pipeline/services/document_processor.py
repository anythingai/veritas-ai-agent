import os
from typing import List, Dict, Any
import structlog
import aiofiles
import docx
import pypdf
from bs4 import BeautifulSoup

logger = structlog.get_logger()

class DocumentProcessor:
    """Handles document parsing and text extraction"""
    
    def __init__(self):
        self.supported_mime_types = {
            'application/pdf': self._process_pdf,
            'text/plain': self._process_text,
            'text/html': self._process_html,
            'application/msword': self._process_doc,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.ms-excel': self._process_excel,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_excel,
            'text/csv': self._process_csv
        }
    
    async def process_document(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Process a document and extract text content"""
        try:
            logger.info("Processing document", file_path=file_path, mime_type=mime_type)
            
            if mime_type not in self.supported_mime_types:
                raise ValueError(f"Unsupported MIME type: {mime_type}")
            
            # Extract text content
            processor = self.supported_mime_types[mime_type]
            content = await processor(file_path)
            
            # Chunk the content
            chunks = self._chunk_content(content)
            
            # Extract metadata
            metadata = self._extract_metadata(file_path, mime_type)
            
            logger.info("Document processing completed", 
                       file_path=file_path,
                       content_length=len(content),
                       chunks_count=len(chunks))
            
            return {
                'content': content,
                'chunks': chunks,
                'metadata': metadata,
                'mime_type': mime_type
            }
            
        except Exception as e:
            logger.error("Document processing failed", 
                        file_path=file_path, 
                        error=str(e))
            raise
    
    async def _process_pdf(self, file_path: str) -> str:
        """Process PDF document using pypdf"""
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content.append(page.extract_text())
            return '\n'.join(text_content)
        except Exception as e:
            logger.error("PDF processing failed", file_path=file_path, error=str(e))
            raise
    
    async def _process_text(self, file_path: str) -> str:
        """Process plain text document"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            async with aiofiles.open(file_path, 'r', encoding='latin-1') as f:
                return await f.read()
    
    async def _process_html(self, file_path: str) -> str:
        """Process HTML document using BeautifulSoup"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            soup = BeautifulSoup(content, 'html.parser')
            return soup.get_text()
        except Exception as e:
            logger.error("HTML processing failed", file_path=file_path, error=str(e))
            raise
    
    async def _process_doc(self, file_path: str) -> str:
        """Process Microsoft Word document (legacy format)"""
        try:
            # For legacy .doc files, we'll use a simple text extraction
            # In production, you might want to use antiword or similar
            async with aiofiles.open(file_path, 'rb') as f:
                content = await f.read()
            # Basic text extraction - this is a simplified approach
            # In production, consider using antiword or similar tools
            return str(content)[:1000] + " [Legacy DOC format - limited text extraction]"
        except Exception as e:
            logger.error("DOC processing failed", file_path=file_path, error=str(e))
            raise
    
    async def _process_docx(self, file_path: str) -> str:
        """Process Microsoft Word document (DOCX) using python-docx"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            return '\n'.join(text_content)
        except Exception as e:
            logger.error("DOCX processing failed", file_path=file_path, error=str(e))
            raise
    
    async def _process_excel(self, file_path: str) -> str:
        """Process Excel document"""
        try:
            # For Excel files, we'll use a simple approach
            # In production, consider using openpyxl or pandas
            async with aiofiles.open(file_path, 'rb') as f:
                content = await f.read()
            return str(content)[:1000] + " [Excel format - limited text extraction]"
        except Exception as e:
            logger.error("Excel processing failed", file_path=file_path, error=str(e))
            raise
    
    async def _process_csv(self, file_path: str) -> str:
        """Process CSV document"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
            return content
        except UnicodeDecodeError:
            async with aiofiles.open(file_path, 'r', encoding='latin-1') as f:
                return await f.read()
    
    def _chunk_content(self, content: str, max_chunk_size: int = 500, overlap: int = 20) -> List[str]:
        """Split content into overlapping chunks"""
        if not content:
            return []
        
        # Simple tokenization (split by whitespace)
        tokens = content.split()
        chunks = []
        
        for i in range(0, len(tokens), max_chunk_size - overlap):
            chunk = ' '.join(tokens[i:i + max_chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks
    
    def _extract_metadata(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        try:
            # Basic metadata extraction
            file_stats = os.stat(file_path)
            metadata = {
                'file_size': file_stats.st_size,
                'mime_type': mime_type,
                'creation_time': file_stats.st_ctime,
                'modification_time': file_stats.st_mtime
            }
            
            # Try to extract additional metadata based on file type
            if mime_type == 'application/pdf':
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = pypdf.PdfReader(file)
                        if pdf_reader.metadata:
                            metadata.update({
                                'title': pdf_reader.metadata.get('/Title', ''),
                                'author': pdf_reader.metadata.get('/Author', ''),
                                'subject': pdf_reader.metadata.get('/Subject', ''),
                                'page_count': len(pdf_reader.pages)
                            })
                except Exception as e:
                    logger.warning("Failed to extract PDF metadata", error=str(e))
            
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                try:
                    doc = docx.Document(file_path)
                    metadata.update({
                        'page_count': len(doc.sections),
                        'paragraph_count': len(doc.paragraphs)
                    })
                except Exception as e:
                    logger.warning("Failed to extract DOCX metadata", error=str(e))
            
            return metadata
            
        except Exception as e:
            logger.warning("Failed to extract metadata", file_path=file_path, error=str(e))
            return {
                'file_size': os.path.getsize(file_path),
                'mime_type': mime_type
            }
    
    def get_supported_mime_types(self) -> List[str]:
        """Get list of supported MIME types"""
        return list(self.supported_mime_types.keys())
    
    def is_supported(self, mime_type: str) -> bool:
        """Check if MIME type is supported"""
        return mime_type in self.supported_mime_types 